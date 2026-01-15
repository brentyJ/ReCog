"""
Microsoft Word (.docx) parser.

Extracts text, tables, headers/footers, comments, and metadata from
Word documents using python-docx library.

Copyright (c) 2025 Brent Lefebure / EhkoLabs
Licensed under AGPLv3
"""

import re
from pathlib import Path
from typing import List, Optional, Dict, Any
from datetime import datetime

from .base import BaseParser
from ..types import ParsedContent


class DocxParser(BaseParser):
    """Parse Microsoft Word .docx files."""

    PARSER_METADATA = {
        "file_type": "Microsoft Word Document",
        "extensions": [".docx"],
        "cypher_context": {
            "description": "Word document with formatted text, tables, and metadata",
            "extractable": [
                "Document text and structure",
                "Tables and lists",
                "Headers and footers",
                "Author and revision history",
                "Comments and annotations"
            ],
            "suggestions": [
                "Author metadata may reveal document origin",
                "Tables can contain structured data points",
                "Check revision count for editing history"
            ]
        }
    }

    def get_extensions(self) -> List[str]:
        return [".docx"]

    def can_parse(self, path: Path) -> bool:
        if path.suffix.lower() != ".docx":
            return False
        # Verify it's a valid ZIP (DOCX is a ZIP archive)
        try:
            import zipfile
            return zipfile.is_zipfile(path)
        except Exception:
            return False

    def get_file_type(self) -> str:
        return "docx"

    def parse(self, path: Path) -> ParsedContent:
        """
        Parse DOCX file and extract all content.

        Args:
            path: Path to .docx file

        Returns:
            ParsedContent with text, tables, and metadata
        """
        try:
            from docx import Document
        except ImportError:
            return ParsedContent(
                text="[DOCX parsing requires python-docx: pip install python-docx]",
                title=path.stem,
                metadata={"error": "python-docx not installed"}
            )

        try:
            doc = Document(str(path))
        except Exception as e:
            return ParsedContent(
                text=f"[Error opening DOCX file: {e}]",
                title=path.stem,
                metadata={"error": "open_failed", "details": str(e)}
            )

        sections = []

        # Extract headers (from first section)
        header_text = self._extract_headers(doc)
        if header_text:
            sections.append(f"[HEADER]\n{header_text}")

        # Extract main body paragraphs
        body_parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                # Check if it's a heading
                if para.style and para.style.name.startswith('Heading'):
                    level = self._get_heading_level(para.style.name)
                    prefix = "#" * level + " " if level else ""
                    body_parts.append(f"\n{prefix}{text}\n")
                else:
                    body_parts.append(text)

        if body_parts:
            sections.append("\n".join(body_parts))

        # Extract tables
        tables_text = self._extract_tables(doc)
        if tables_text:
            sections.append(f"\n[TABLES]\n{tables_text}")

        # Extract footers (from first section)
        footer_text = self._extract_footers(doc)
        if footer_text:
            sections.append(f"\n[FOOTER]\n{footer_text}")

        # Extract comments if available
        comments_text = self._extract_comments(path)
        if comments_text:
            sections.append(f"\n[COMMENTS]\n{comments_text}")

        # Combine all sections
        full_text = "\n\n".join(sections)

        # Extract metadata
        metadata = self._extract_metadata(doc, path)

        # Determine title
        title = None
        if doc.core_properties.title:
            title = doc.core_properties.title
        elif body_parts:
            # Use first non-empty paragraph as title
            title = body_parts[0][:100] if body_parts[0] else path.stem
        else:
            title = path.stem

        return ParsedContent(
            text=full_text,
            title=title,
            author=doc.core_properties.author,
            date=self._format_date(doc.core_properties.created),
            metadata=metadata
        )

    def _extract_headers(self, doc) -> str:
        """Extract header text from document."""
        headers = []
        try:
            for section in doc.sections:
                header = section.header
                if header:
                    for para in header.paragraphs:
                        text = para.text.strip()
                        if text:
                            headers.append(text)
        except Exception:
            pass
        return "\n".join(headers) if headers else ""

    def _extract_footers(self, doc) -> str:
        """Extract footer text from document."""
        footers = []
        try:
            for section in doc.sections:
                footer = section.footer
                if footer:
                    for para in footer.paragraphs:
                        text = para.text.strip()
                        if text:
                            footers.append(text)
        except Exception:
            pass
        return "\n".join(footers) if footers else ""

    def _extract_tables(self, doc) -> str:
        """Extract tables as text representation."""
        tables = []
        for i, table in enumerate(doc.tables, 1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                table_text = f"Table {i}:\n" + "\n".join(rows)
                tables.append(table_text)
        return "\n\n".join(tables)

    def _extract_comments(self, path: Path) -> str:
        """Extract comments from DOCX (requires parsing XML directly)."""
        comments = []
        try:
            import zipfile
            from xml.etree import ElementTree as ET

            with zipfile.ZipFile(path, 'r') as zf:
                if 'word/comments.xml' in zf.namelist():
                    with zf.open('word/comments.xml') as f:
                        tree = ET.parse(f)
                        root = tree.getroot()

                        # Define namespace
                        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}

                        for comment in root.findall('.//w:comment', ns):
                            author = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author', 'Unknown')
                            date = comment.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date', '')

                            # Extract comment text
                            text_parts = []
                            for t in comment.findall('.//w:t', ns):
                                if t.text:
                                    text_parts.append(t.text)

                            if text_parts:
                                comment_text = "".join(text_parts)
                                comments.append(f"[{author}] {comment_text}")
        except Exception:
            pass

        return "\n".join(comments)

    def _extract_metadata(self, doc, path: Path) -> Dict[str, Any]:
        """Extract document metadata."""
        props = doc.core_properties
        metadata = {
            "format": "docx",
            "paragraph_count": len(doc.paragraphs),
            "table_count": len(doc.tables),
        }

        # Core properties
        if props.author:
            metadata["author"] = props.author
        if props.title:
            metadata["title"] = props.title
        if props.subject:
            metadata["subject"] = props.subject
        if props.keywords:
            metadata["keywords"] = props.keywords
        if props.created:
            metadata["created"] = self._format_date(props.created)
        if props.modified:
            metadata["modified"] = self._format_date(props.modified)
        if props.last_modified_by:
            metadata["last_modified_by"] = props.last_modified_by
        if props.revision is not None:
            metadata["revision_count"] = props.revision
        if props.category:
            metadata["category"] = props.category

        # Word count (approximate)
        word_count = sum(len(p.text.split()) for p in doc.paragraphs)
        metadata["word_count"] = word_count

        return metadata

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        match = re.search(r'Heading\s*(\d+)', style_name)
        if match:
            return int(match.group(1))
        return 0

    def _format_date(self, dt) -> Optional[str]:
        """Format datetime to ISO date string."""
        if dt is None:
            return None
        if isinstance(dt, datetime):
            return dt.date().isoformat()
        return str(dt)


__all__ = ["DocxParser"]
